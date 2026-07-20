import hashlib
import io
import json
import logging
import re
import zipfile
from datetime import datetime, timedelta, timezone
from urllib.parse import urlsplit

from fastapi import BackgroundTasks, Depends, FastAPI, HTTPException, Query, Request, UploadFile, File, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, RedirectResponse, Response, StreamingResponse
import os
from jose import JWTError, jwt
from pypdf import PdfReader
from docx import Document as DocxDocument
from sqlalchemy import inspect, text
from sqlalchemy.orm import Session

from app.auth import (
    ALGORITHM,
    SECRET_KEY,
    create_access_token,
    get_current_user,
    hash_password,
    verify_password,
)
from app.database import Base, SessionLocal, engine, get_db
from app.models import (
    ChatChannel,
    ChatMessage,
    ChatSession,
    DataSourceConnection,
    Document,
    GeneratedReport,
    Integration,
    Settings,
    TelegramChatLink,
    User,
)
from app.crypto import decrypt_str, encrypt_str
from app.telegram import (
    generate_webhook_secret,
    handle_telegram_update,
    telegram_delete_webhook,
    telegram_get_me,
    telegram_get_webhook_info,
    telegram_set_webhook,
)
from app.google_drive import (
    encode_connection_tokens,
    exchange_code_for_credentials,
    fetch_account_email,
    get_authorization_url,
    get_picker_access_token,
)
from app.google_drive import sync_connection as sync_google_drive_connection
from app.rag import delete_document_chunks, delete_user_collection, process_document
from app.query import generate_answer_stream, generate_report
from app.scheduler import PriorityWorkerPool, PrioritySlotGate, plan_priority
from app.schemas import (
    ChatChannelIn,
    ChatChannelOut,
    ChatMessageOut,
    ChatSessionOut,
    ChatSessionRenameIn,
    DataSourceConnectionIn,
    DataSourceConnectionOut,
    DataSourceSyncStartedOut,
    DbStatusOut,
    DocumentOut,
    GeneratedReportOut,
    GoogleDriveAuthorizeOut,
    GoogleDriveFoldersIn,
    GoogleDrivePickerTokenOut,
    IntegrationIn,
    IntegrationOut,
    PlanOut,
    ReportExportIn,
    ReportOut,
    SettingsIn,
    SettingsOut,
    StatsOut,
    SubscribeIn,
    TokenOut,
    UserCreate,
    UserLogin,
    UserOut,
    UserUpdate,
)

IMAGE_EXTENSIONS = (".jpg", ".jpeg", ".png")


def extract_image_text(file_path):
    try:
        import pytesseract
        from PIL import Image

        return pytesseract.image_to_string(Image.open(file_path)).strip()
    except Exception:
        return ""


def extract_document_text(file_path, filename):
    """Returns (text, file_type, page_count) for a file already on disk, or
    (None, None, None) if the extension isn't supported."""
    text = ""
    file_type = None
    page_count = None
    filename_lower = filename.lower()

    if filename_lower.endswith(".pdf"):
        file_type = "pdf"
        reader = PdfReader(file_path)
        page_count = len(reader.pages)

        for page in reader.pages:
            extracted_text = page.extract_text()
            if extracted_text:
                text += extracted_text + "\n"

    elif filename_lower.endswith(".docx"):
        file_type = "docx"
        doc = DocxDocument(file_path)

        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

    elif filename_lower.endswith(IMAGE_EXTENSIONS):
        file_type = "image"
        text = extract_image_text(file_path)

    else:
        return None, None, None

    return text, file_type, page_count


logger = logging.getLogger("uvicorn.error")

try:
    Base.metadata.create_all(bind=engine)
except Exception:
    logger.exception("Base.metadata.create_all failed; continuing startup")


def ensure_chat_message_session_column():
    """create_all only creates missing tables, not new columns on tables that
    already existed -- chat_messages predates the session_id column, so add it
    by hand if it isn't there yet. Idempotent: checked via information_schema
    every startup, but only ALTERs once."""
    inspector = inspect(engine)
    columns = {col["name"] for col in inspector.get_columns("chat_messages")}
    if "session_id" in columns:
        return

    with engine.connect() as conn:
        conn.execute(text("ALTER TABLE chat_messages ADD COLUMN session_id INTEGER NULL"))
        conn.execute(
            text("ALTER TABLE chat_messages ADD INDEX ix_chat_messages_session_id (session_id)")
        )
        conn.commit()


try:
    ensure_chat_message_session_column()
except Exception:
    logger.exception("ensure_chat_message_session_column failed; continuing startup")


def ensure_data_source_connection_columns():
    """data_source_connections predates config/status/last_synced_at (added for
    the Google Drive connector) -- add them by hand if missing, same pattern as
    ensure_chat_message_session_column above."""
    inspector = inspect(engine)
    existing_columns = {col["name"]: col for col in inspector.get_columns("data_source_connections")}
    columns = set(existing_columns)
    credential_col = existing_columns.get("credential")

    with engine.connect() as conn:
        if credential_col is not None and "TEXT" not in str(credential_col["type"]).upper():
            # Service account JSON keys are far longer than the original
            # VARCHAR(255) credential column allowed.
            conn.execute(text("ALTER TABLE data_source_connections MODIFY credential TEXT NULL"))
        if "config" not in columns:
            conn.execute(text("ALTER TABLE data_source_connections ADD COLUMN config TEXT NULL"))
        if "status" not in columns:
            conn.execute(text("ALTER TABLE data_source_connections ADD COLUMN status VARCHAR(20) NULL"))
        if "status_message" not in columns:
            conn.execute(text("ALTER TABLE data_source_connections ADD COLUMN status_message TEXT NULL"))
        if "last_synced_at" not in columns:
            conn.execute(text("ALTER TABLE data_source_connections ADD COLUMN last_synced_at TIMESTAMP NULL"))
        conn.commit()


try:
    ensure_data_source_connection_columns()
except Exception:
    logger.exception("ensure_data_source_connection_columns failed; continuing startup")


def ensure_document_source_columns():
    """documents predates source_type/data_source_id/external_id (added so a
    Google Drive sync can tell an already-imported file from one removed at
    the source) -- add them by hand if missing."""
    inspector = inspect(engine)
    columns = {col["name"] for col in inspector.get_columns("documents")}
    if "external_id" in columns:
        return

    with engine.connect() as conn:
        conn.execute(text("ALTER TABLE documents ADD COLUMN source_type VARCHAR(30) NULL"))
        conn.execute(text("ALTER TABLE documents ADD COLUMN data_source_id INTEGER NULL"))
        conn.execute(text("ALTER TABLE documents ADD COLUMN external_id VARCHAR(255) NULL"))
        conn.execute(text("ALTER TABLE documents ADD INDEX ix_documents_data_source_id (data_source_id)"))
        conn.execute(text("ALTER TABLE documents ADD INDEX ix_documents_external_id (external_id)"))
        conn.commit()


try:
    ensure_document_source_columns()
except Exception:
    logger.exception("ensure_document_source_columns failed; continuing startup")


def ensure_chat_channel_columns():
    """chat_channels predates bot_username/webhook_secret/updated_at (added for
    the Telegram bot connector) -- add them by hand if missing, same pattern
    as ensure_data_source_connection_columns above. Also widen bot_token to
    TEXT since it now stores an encrypted (longer) value instead of the raw
    token."""
    inspector = inspect(engine)
    existing_columns = {col["name"]: col for col in inspector.get_columns("chat_channels")}
    columns = set(existing_columns)
    bot_token_col = existing_columns.get("bot_token")

    with engine.connect() as conn:
        if bot_token_col is not None and "TEXT" not in str(bot_token_col["type"]).upper():
            conn.execute(text("ALTER TABLE chat_channels MODIFY bot_token TEXT NULL"))
        if "bot_username" not in columns:
            conn.execute(text("ALTER TABLE chat_channels ADD COLUMN bot_username VARCHAR(150) NULL"))
        if "webhook_secret" not in columns:
            conn.execute(text("ALTER TABLE chat_channels ADD COLUMN webhook_secret VARCHAR(64) NULL"))
            conn.execute(text("ALTER TABLE chat_channels ADD UNIQUE INDEX ix_chat_channels_webhook_secret (webhook_secret)"))
        if "updated_at" not in columns:
            conn.execute(text("ALTER TABLE chat_channels ADD COLUMN updated_at TIMESTAMP NULL"))
        conn.commit()


try:
    ensure_chat_channel_columns()
except Exception:
    logger.exception("ensure_chat_channel_columns failed; continuing startup")


def migrate_orphan_chat_messages():
    """One-time backfill for chat_messages saved before sessions existed.
    Idempotent: once a message has a session_id, it's never picked up again."""
    db = SessionLocal()
    try:
        user_ids = [
            row[0]
            for row in db.query(ChatMessage.user_id)
            .filter(ChatMessage.session_id.is_(None))
            .distinct()
            .all()
        ]
        for user_id in user_ids:
            orphans = (
                db.query(ChatMessage)
                .filter(ChatMessage.user_id == user_id, ChatMessage.session_id.is_(None))
                .order_by(ChatMessage.created_at.asc())
                .all()
            )
            if not orphans:
                continue

            title = (orphans[0].question or "Previous conversation")[:60]
            session = ChatSession(user_id=user_id, title=title)
            db.add(session)
            db.flush()

            for message in orphans:
                message.session_id = session.id

            db.commit()
    finally:
        db.close()


try:
    migrate_orphan_chat_messages()
except Exception:
    logger.exception("migrate_orphan_chat_messages failed; continuing startup")

app = FastAPI(
    title="Cognexa API",
    version="1.0.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000",
        "https://cognexarag.onrender.com",
        "https://cognexa-api.onrender.com",
        "https://*.onrender.com",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_FOLDER = "app/uploads"

os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Priority scheduling backing the plan pricing claims: paid plans' document
# indexing jobs and chat generations are served ahead of Community's when
# multiple are in flight at once. See app/scheduler.py.
INDEX_POOL = PriorityWorkerPool(num_workers=3)
RETRIEVAL_GATE = PrioritySlotGate(capacity=2)


def index_document_job(document_id, text, filename, user_id, chunk_size, chunk_overlap):
    db = SessionLocal()
    try:
        result = process_document(
            text,
            filename,
            user_id=user_id,
            document_id=document_id,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
        )
        document = db.query(Document).filter(Document.id == document_id).first()
        if document:
            document.chunks = result["chunks"]
            db.commit()
    finally:
        db.close()


def sync_google_drive_connection_job(connection_id, upload_folder, chunk_size, chunk_overlap):
    db = SessionLocal()
    try:
        connection = db.query(DataSourceConnection).filter(DataSourceConnection.id == connection_id).first()
        if not connection:
            return
        try:
            sync_google_drive_connection(
                db,
                connection,
                extract_document_text,
                upload_folder,
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
            )
        except Exception as exc:
            connection.status = "error"
            connection.status_message = f"Sync failed: {exc}"
            db.commit()
    finally:
        db.close()

PLAN_LIMITS = {
    "community": {
        "max_documents": 25,
        "max_storage_bytes": 15 * 1024 * 1024,
        "max_apps": 2,
        "max_chat_channels": 1,
    },
    "pro": {
        "max_documents": 100,
        "max_storage_bytes": 10 * 1024 * 1024 * 1024,
        "max_apps": 10,
        "max_chat_channels": 5,
    },
    "team": {
        "max_documents": None,
        "max_storage_bytes": None,
        "max_apps": None,
        "max_chat_channels": None,
    },
}

BILLING_CYCLE_LENGTH = timedelta(days=365)

# "community" stays the internal plan key (stored on users, compared in code)
# — only this display name changed, to "Free" per product naming.
PLAN_DISPLAY_NAMES = {
    "community": "Free",
    "pro": "Pro",
    "team": "Unlimited",
}

OPENROUTER_FREE_MODELS = {
    "openai/gpt-oss-20b:free",
    "google/gemma-4-26b-a4b-it:free",
    "cohere/north-mini-code:free",
    "tencent/hy3:free",
}

INTEGRATION_LIMITS = {
    "community": 1,
    "pro": 3,
    "team": None,
}

# Only the Community plan meters AI-provider usage; Pro/Unlimited use their own
# API usage/keys and are never capped here.
COMMUNITY_MONTHLY_AI_CREDITS = 50
CREDIT_PERIOD = timedelta(days=30)


def get_remaining_community_credits(user, db) -> int:
    now = datetime.utcnow()
    period_start = user.ai_credits_period_start

    if period_start is None or now - period_start >= CREDIT_PERIOD:
        user.ai_credits_used = 0
        user.ai_credits_period_start = now
        db.commit()

    return max(0, COMMUNITY_MONTHLY_AI_CREDITS - (user.ai_credits_used or 0))


def consume_community_credit(user, db):
    user.ai_credits_used = (user.ai_credits_used or 0) + 1
    db.commit()


def is_allowed_community_integration(provider_name, model, base_url):
    return (
        provider_name == "OpenRouter"
        and model in OPENROUTER_FREE_MODELS
        and not base_url
    )


def get_plan_usage(db: Session, user_id: int):
    documents = db.query(Document).filter(Document.user_id == user_id).all()
    document_count = len(documents)
    storage_bytes = sum(doc.size_bytes or 0 for doc in documents)
    return document_count, storage_bytes


def get_or_create_settings(db: Session, user_id: int) -> Settings:
    settings = db.query(Settings).filter(Settings.user_id == user_id).first()

    if not settings:
        settings = Settings(user_id=user_id)
        db.add(settings)
        db.commit()
        db.refresh(settings)

    return settings


@app.get("/")
def home():
    return {
        "message": "Cognexa API Running"
    }


@app.get("/health")
def health():
    return {
        "status": "online",
        "message": "Backend connected successfully"
    }


@app.get("/health/db", response_model=DbStatusOut)
def health_db(db: Session = Depends(get_db)):
    try:
        db.execute(text("SELECT 1"))
        return {"connected": True, "message": "Database connected successfully"}
    except Exception as e:
        return {"connected": False, "message": str(e)}


@app.post("/auth/register", response_model=TokenOut)
def register(user_in: UserCreate, db: Session = Depends(get_db)):
    existing = db.query(User).filter(User.email == user_in.email).first()

    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")

    user = User(
        name=user_in.name,
        email=user_in.email,
        password=hash_password(user_in.password),
    )

    db.add(user)
    db.commit()
    db.refresh(user)

    token = create_access_token({"sub": str(user.id)})

    return {"access_token": token, "user": user}


@app.post("/auth/login", response_model=TokenOut)
def login(credentials: UserLogin, db: Session = Depends(get_db)):
    user = db.query(User).filter(User.email == credentials.email).first()

    if not user or not verify_password(credentials.password, user.password):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid email or password",
        )

    token = create_access_token({"sub": str(user.id)})

    return {"access_token": token, "user": user}


@app.get("/auth/me", response_model=UserOut)
def me(current_user: User = Depends(get_current_user)):
    return current_user


@app.patch("/auth/me", response_model=UserOut)
def update_me(
    user_in: UserUpdate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    current_user.name = user_in.name

    if user_in.password:
        current_user.password = hash_password(user_in.password)

    db.commit()
    db.refresh(current_user)

    return current_user


@app.get("/billing/plan", response_model=PlanOut)
def get_billing_plan(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    plan = current_user.plan or "community"
    limits = PLAN_LIMITS.get(plan, PLAN_LIMITS["community"])
    document_count, storage_bytes = get_plan_usage(db, current_user.id)
    cycle_start = current_user.created_at or datetime.utcnow()

    return {
        "plan": plan,
        "max_documents": limits["max_documents"],
        "max_storage_bytes": limits["max_storage_bytes"],
        "document_count": document_count,
        "storage_bytes": storage_bytes,
        "max_ai_credits": COMMUNITY_MONTHLY_AI_CREDITS if plan == "community" else None,
        "ai_credits_remaining": (
            get_remaining_community_credits(current_user, db) if plan == "community" else None
        ),
        "max_apps": limits["max_apps"],
        "apps_connected": db.query(DataSourceConnection).filter(
            DataSourceConnection.user_id == current_user.id
        ).count(),
        "max_chat_channels": limits["max_chat_channels"],
        "chat_channels_connected": db.query(ChatChannel).filter(
            ChatChannel.user_id == current_user.id
        ).count(),
        "billing_cycle_start": cycle_start,
        "billing_cycle_end": cycle_start + BILLING_CYCLE_LENGTH,
    }


@app.post("/billing/subscribe", response_model=PlanOut)
def subscribe_plan(
    subscribe_in: SubscribeIn,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if subscribe_in.plan not in PLAN_LIMITS:
        raise HTTPException(status_code=400, detail="Unknown plan")

    # Demo checkout — no real payment processor wired up. Any submitted card
    # details are accepted and discarded; the plan is applied immediately.
    current_user.plan = subscribe_in.plan
    db.commit()
    db.refresh(current_user)

    limits = PLAN_LIMITS[subscribe_in.plan]
    document_count, storage_bytes = get_plan_usage(db, current_user.id)
    plan = current_user.plan
    cycle_start = current_user.created_at or datetime.utcnow()

    return {
        "plan": plan,
        "max_documents": limits["max_documents"],
        "max_storage_bytes": limits["max_storage_bytes"],
        "document_count": document_count,
        "storage_bytes": storage_bytes,
        "max_ai_credits": COMMUNITY_MONTHLY_AI_CREDITS if plan == "community" else None,
        "ai_credits_remaining": (
            get_remaining_community_credits(current_user, db) if plan == "community" else None
        ),
        "max_apps": limits["max_apps"],
        "apps_connected": db.query(DataSourceConnection).filter(
            DataSourceConnection.user_id == current_user.id
        ).count(),
        "max_chat_channels": limits["max_chat_channels"],
        "chat_channels_connected": db.query(ChatChannel).filter(
            ChatChannel.user_id == current_user.id
        ).count(),
        "billing_cycle_start": cycle_start,
        "billing_cycle_end": cycle_start + BILLING_CYCLE_LENGTH,
    }


@app.post("/upload")
async def upload_file(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    plan = current_user.plan or "community"
    limits = PLAN_LIMITS.get(plan, PLAN_LIMITS["community"])
    plan_display_name = PLAN_DISPLAY_NAMES.get(plan, plan.capitalize())
    document_count, storage_bytes = get_plan_usage(db, current_user.id)
    upgrade_hint = "Upgrade your plan for unlimited documents." if plan != "team" else ""

    if limits["max_documents"] is not None and document_count >= limits["max_documents"]:
        raise HTTPException(
            status_code=402,
            detail=(
                f"{plan_display_name} plan limit reached ({limits['max_documents']} documents). "
                f"{upgrade_hint}"
            ),
        )

    file_path = os.path.join(
        UPLOAD_FOLDER,
        file.filename
    )

    contents = await file.read()
    settings = get_or_create_settings(db, current_user.id)

    if limits["max_storage_bytes"] is not None and (
        storage_bytes + len(contents) > limits["max_storage_bytes"]
    ):
        storage_upgrade_hint = "Upgrade your plan for more storage." if plan != "team" else ""
        raise HTTPException(
            status_code=402,
            detail=(
                f"{plan_display_name} plan storage limit reached "
                f"({limits['max_storage_bytes'] // (1024 * 1024)} MB). "
                f"{storage_upgrade_hint}"
            ),
        )

    content_hash = hashlib.sha256(contents).hexdigest()

    if settings.duplicate_detection:
        existing = (
            db.query(Document)
            .filter(Document.user_id == current_user.id, Document.content_hash == content_hash)
            .first()
        )
        if existing:
            raise HTTPException(
                status_code=409,
                detail=f'This file is identical to the already-uploaded "{existing.filename}".',
            )

    with open(file_path, "wb") as buffer:
        buffer.write(contents)

    text, file_type, page_count = extract_document_text(file_path, file.filename)
    if file_type is None:
        return {
            "error": "Only PDF, DOCX, JPG and PNG files are supported"
        }

    preview_text = text[:1000] if text.strip() else "Image uploaded (no text detected)."

    document = Document(
        user_id=current_user.id,
        filename=file.filename,
        file_path=file_path,
        file_type=file_type,
        chunks=0,
        preview=preview_text,
        size_bytes=os.path.getsize(file_path),
        page_count=page_count,
        content_hash=content_hash,
    )

    db.add(document)
    db.commit()
    db.refresh(document)

    INDEX_POOL.submit(
        plan_priority(plan),
        index_document_job,
        document.id,
        text,
        file.filename,
        current_user.id,
        settings.chunk_size,
        settings.chunk_overlap,
    )

    return {
        "id": document.id,
        "filename": file.filename,
        "characters": len(text),
        "chunks_saved": 0,
        "processing": True,
        "preview": preview_text
    }


@app.get("/documents", response_model=list[DocumentOut])
def list_documents(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    return (
        db.query(Document)
        .filter(Document.user_id == current_user.id)
        .order_by(Document.created_at.desc())
        .all()
    )


@app.get("/documents/{document_id}", response_model=DocumentOut)
def get_document(
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    document = (
        db.query(Document)
        .filter(Document.id == document_id, Document.user_id == current_user.id)
        .first()
    )
    if not document:
        raise HTTPException(status_code=404, detail="Document not found.")

    return document


@app.get("/documents/{document_id}/download")
def download_document(
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    document = (
        db.query(Document)
        .filter(Document.id == document_id, Document.user_id == current_user.id)
        .first()
    )

    if not document or not document.file_path or not os.path.exists(document.file_path):
        raise HTTPException(status_code=404, detail="Document not found")

    return FileResponse(document.file_path, filename=document.filename)


@app.post("/documents/{document_id}/reindex")
def reindex_document(
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    document = (
        db.query(Document)
        .filter(Document.id == document_id, Document.user_id == current_user.id)
        .first()
    )

    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    if not document.file_path or not os.path.exists(document.file_path):
        raise HTTPException(
            status_code=404,
            detail="The original file is no longer available, so it can't be re-indexed.",
        )

    text, file_type, _ = extract_document_text(document.file_path, document.filename)
    if file_type is None:
        raise HTTPException(status_code=400, detail="Unsupported file type.")

    settings = get_or_create_settings(db, current_user.id)
    plan = current_user.plan or "community"

    delete_document_chunks(current_user.id, document.id)
    document.chunks = 0
    db.commit()

    INDEX_POOL.submit(
        plan_priority(plan),
        index_document_job,
        document.id,
        text,
        document.filename,
        current_user.id,
        settings.chunk_size,
        settings.chunk_overlap,
    )

    return {"queued": True}


@app.delete("/documents/{document_id}")
def delete_document(
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    document = (
        db.query(Document)
        .filter(Document.id == document_id, Document.user_id == current_user.id)
        .first()
    )

    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    if document.file_path and os.path.exists(document.file_path):
        os.remove(document.file_path)

    delete_document_chunks(current_user.id, document.id)

    db.delete(document)
    db.commit()

    return {"deleted": document_id}


@app.delete("/documents")
def delete_all_documents(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    documents = db.query(Document).filter(Document.user_id == current_user.id).all()

    for document in documents:
        if document.file_path and os.path.exists(document.file_path):
            os.remove(document.file_path)
        delete_document_chunks(current_user.id, document.id)
        db.delete(document)

    db.commit()

    return {"deleted": len(documents)}


DOCUMENT_LISTING_PATTERN = re.compile(
    r"\b(list|what|which|how many)\b.*\b(document|documents|file|files|upload|uploads|knowledge base)\b",
    re.IGNORECASE,
)


def is_document_listing_question(question: str) -> bool:
    return bool(DOCUMENT_LISTING_PATTERN.search(question))


@app.post("/ask")
async def ask_question(
    question: str,
    session_id: int = Query(...),
    document_ids: list[int] | None = Query(default=None),
    source: str = Query(default="local"),
    integration_id: int | None = Query(default=None),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    session = (
        db.query(ChatSession)
        .filter(ChatSession.id == session_id, ChatSession.user_id == current_user.id)
        .first()
    )
    if not session:
        raise HTTPException(status_code=404, detail="Chat session not found.")

    settings = get_or_create_settings(db, current_user.id)

    if document_ids:
        owned_count = (
            db.query(Document)
            .filter(Document.id.in_(document_ids), Document.user_id == current_user.id)
            .count()
        )
        if owned_count != len(document_ids):
            raise HTTPException(status_code=404, detail="One or more documents not found")

    plan = current_user.plan or "community"

    def resolve_integration(integration):
        if not integration or not integration.api_key:
            return None
        if plan == "community" and not is_allowed_community_integration(
            integration.provider_name, integration.model, integration.base_url
        ):
            return None
        return integration

    selected_integration = None
    if source == "integration" and integration_id is not None:
        requested_integration = (
            db.query(Integration)
            .filter(Integration.id == integration_id, Integration.user_id == current_user.id)
            .first()
        )
        if not requested_integration:
            raise HTTPException(status_code=404, detail="Integration not found.")
        selected_integration = resolve_integration(requested_integration)
        if not selected_integration:
            if not requested_integration.api_key:
                raise HTTPException(
                    status_code=400,
                    detail=(
                        f"The {requested_integration.provider_name} integration has no API key "
                        "saved. Fix it under Settings > Integrations."
                    ),
                )
            raise HTTPException(
                status_code=402,
                detail=(
                    f"The {requested_integration.provider_name} integration "
                    f"({requested_integration.model or 'no model'}) isn't allowed on the "
                    "Community plan. Pick one of the free OpenRouter models under "
                    "Settings > Integrations, or upgrade your plan."
                ),
            )
    elif source == "auto":
        auto_integration = (
            db.query(Integration)
            .filter(Integration.user_id == current_user.id)
            .order_by(Integration.created_at.desc())
            .first()
        )
        selected_integration = resolve_integration(auto_integration)
        if (
            auto_integration
            and not selected_integration
            and auto_integration.api_key
            and plan == "community"
        ):
            raise HTTPException(
                status_code=402,
                detail=(
                    f"Your saved {auto_integration.provider_name} integration "
                    f"({auto_integration.model or 'no model'}) isn't allowed on the "
                    "Community plan. Pick one of the free OpenRouter models under "
                    "Settings > Integrations, or upgrade your plan."
                ),
            )

    # Community's OpenRouter integration is metered — once the monthly credit
    # allowance is used up, silently fall back to the local model instead of
    # calling OpenRouter. Pro/Unlimited use their own API usage and are never
    # capped here.
    if selected_integration and plan == "community":
        if get_remaining_community_credits(current_user, db) <= 0:
            selected_integration = None

    def save_message(answer_text, sources_list):
        message = ChatMessage(
            user_id=current_user.id,
            session_id=session.id,
            question=question,
            answer=answer_text,
            sources=",".join(sources_list),
        )
        db.add(message)

        if session.title == "New Chat" or not session.title:
            session.title = question[:60]
        session.updated_at = datetime.utcnow()

        db.commit()

    def event_stream():
        full_answer = []
        sources = []

        try:
            # "List/what/how many documents..." is a metadata question, not a
            # content question — answering it via embedding similarity search
            # is unreliable (it only surfaces documents whose chunks happen to
            # rank high for that phrasing, which skews by document type/length
            # and can drop documents from the answer entirely). The document
            # list already exists in the database, so answer it directly.
            if not document_ids and is_document_listing_question(question):
                owned_documents = (
                    db.query(Document)
                    .filter(Document.user_id == current_user.id)
                    .order_by(Document.created_at.asc())
                    .all()
                )
                sources = sorted({doc.filename for doc in owned_documents})

                yield f"data: {json.dumps({'type': 'sources', 'sources': sources})}\n\n"

                if not owned_documents:
                    answer_text = "You haven't uploaded any documents yet."
                else:
                    answer_text = "You've uploaded:\n" + "\n".join(
                        f"- {doc.filename}" for doc in owned_documents
                    )

                full_answer.append(answer_text)
                yield f"data: {json.dumps({'type': 'token', 'content': answer_text})}\n\n"

                save_message("".join(full_answer), sources)

                yield f"data: {json.dumps({'type': 'done', 'sources': sources})}\n\n"
                return

            if selected_integration and plan == "community":
                consume_community_credit(current_user, db)

            with RETRIEVAL_GATE(plan_priority(plan)):
                for event in generate_answer_stream(
                    question,
                    user_id=current_user.id,
                    document_ids=document_ids,
                    ollama_url=settings.ollama_url,
                    llm_model=settings.llm_model,
                    external_provider=selected_integration.provider_name if selected_integration else None,
                    external_api_key=selected_integration.api_key if selected_integration else None,
                    external_base_url=selected_integration.base_url if selected_integration else None,
                    external_model=selected_integration.model if selected_integration else None,
                ):
                    if event["type"] == "sources":
                        sources = event["sources"]
                    elif event["type"] == "token":
                        full_answer.append(event["content"])

                    yield f"data: {json.dumps(event)}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"
            return

        save_message("".join(full_answer), sources)

        yield f"data: {json.dumps({'type': 'done', 'sources': sources})}\n\n"

    return StreamingResponse(event_stream(), media_type="text/event-stream")


@app.get("/chat/sessions", response_model=list[ChatSessionOut])
def list_chat_sessions(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    sessions = (
        db.query(ChatSession)
        .filter(ChatSession.user_id == current_user.id)
        .order_by(ChatSession.updated_at.desc())
        .all()
    )

    # A session is created eagerly on "New Chat" before any message is sent;
    # if the user navigates away without sending one, it's an empty orphan --
    # hide it here (and clean it up) instead of showing a title-less duplicate.
    # Skip sessions created in the last 2 minutes so we don't delete the one
    # the user just opened and hasn't sent a message in yet.
    cutoff = datetime.now(timezone.utc) - timedelta(minutes=2)
    empty_ids = [
        s.id
        for s in sessions
        if s.title == "New Chat"
        and s.created_at
        and s.created_at.replace(tzinfo=timezone.utc) < cutoff
        and not db.query(ChatMessage.id).filter(ChatMessage.session_id == s.id).first()
    ]
    if empty_ids:
        db.query(ChatSession).filter(ChatSession.id.in_(empty_ids)).delete(synchronize_session=False)
        db.commit()
        sessions = [s for s in sessions if s.id not in empty_ids]

    return sessions


@app.post("/chat/sessions", response_model=ChatSessionOut)
def create_chat_session(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    session = ChatSession(user_id=current_user.id, title="New Chat")
    db.add(session)
    db.commit()
    db.refresh(session)

    return session


@app.patch("/chat/sessions/{session_id}", response_model=ChatSessionOut)
def rename_chat_session(
    session_id: int,
    rename_in: ChatSessionRenameIn,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    session = (
        db.query(ChatSession)
        .filter(ChatSession.id == session_id, ChatSession.user_id == current_user.id)
        .first()
    )
    if not session:
        raise HTTPException(status_code=404, detail="Chat session not found.")

    session.title = rename_in.title.strip()[:60] or "New Chat"
    db.commit()
    db.refresh(session)

    return session


@app.delete("/chat/sessions/{session_id}")
def delete_chat_session(
    session_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    session = (
        db.query(ChatSession)
        .filter(ChatSession.id == session_id, ChatSession.user_id == current_user.id)
        .first()
    )
    if not session:
        raise HTTPException(status_code=404, detail="Chat session not found.")

    db.query(ChatMessage).filter(ChatMessage.session_id == session_id).delete()
    db.delete(session)
    db.commit()

    return {"deleted": session_id}


@app.get("/chat/sessions/{session_id}/messages", response_model=list[ChatMessageOut])
def get_chat_session_messages(
    session_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    session = (
        db.query(ChatSession)
        .filter(ChatSession.id == session_id, ChatSession.user_id == current_user.id)
        .first()
    )
    if not session:
        raise HTTPException(status_code=404, detail="Chat session not found.")

    return (
        db.query(ChatMessage)
        .filter(ChatMessage.session_id == session_id)
        .order_by(ChatMessage.created_at.asc())
        .all()
    )


def resolve_report_integration(db, current_user, plan, source, integration_id):
    def resolve_integration(integration):
        if not integration or not integration.api_key:
            return None
        if plan == "community" and not is_allowed_community_integration(
            integration.provider_name, integration.model, integration.base_url
        ):
            return None
        return integration

    selected_integration = None
    if source == "integration" and integration_id is not None:
        requested_integration = (
            db.query(Integration)
            .filter(Integration.id == integration_id, Integration.user_id == current_user.id)
            .first()
        )
        if not requested_integration:
            raise HTTPException(status_code=404, detail="Integration not found.")
        selected_integration = resolve_integration(requested_integration)
    elif source == "auto":
        auto_integration = (
            db.query(Integration)
            .filter(Integration.user_id == current_user.id)
            .order_by(Integration.created_at.desc())
            .first()
        )
        selected_integration = resolve_integration(auto_integration)

    if selected_integration and plan == "community":
        if get_remaining_community_credits(current_user, db) <= 0:
            selected_integration = None

    return selected_integration


@app.post("/report/session", response_model=ReportOut)
def report_session(
    session_id: int = Query(...),
    source: str = Query(default="auto"),
    integration_id: int | None = Query(default=None),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    chat_session = (
        db.query(ChatSession)
        .filter(ChatSession.id == session_id, ChatSession.user_id == current_user.id)
        .first()
    )
    if not chat_session:
        raise HTTPException(status_code=404, detail="Chat session not found.")

    settings = get_or_create_settings(db, current_user.id)
    plan = current_user.plan or "community"
    selected_integration = resolve_report_integration(db, current_user, plan, source, integration_id)

    messages = (
        db.query(ChatMessage)
        .filter(ChatMessage.session_id == session_id)
        .order_by(ChatMessage.created_at.asc())
        .all()
    )
    qa_pairs = [
        {
            "question": m.question,
            "answer": m.answer,
            "sources": [s for s in (m.sources or "").split(",") if s],
        }
        for m in messages
    ]

    if selected_integration and plan == "community":
        consume_community_credit(current_user, db)

    result = generate_report(
        qa_pairs,
        ollama_url=settings.ollama_url,
        llm_model=settings.llm_model,
        external_provider=selected_integration.provider_name if selected_integration else None,
        external_api_key=selected_integration.api_key if selected_integration else None,
        external_base_url=selected_integration.base_url if selected_integration else None,
        external_model=selected_integration.model if selected_integration else None,
    )

    save_generated_report(
        db,
        current_user,
        session_id=session_id,
        title=chat_session.title,
        report=result["report"],
        sources=result["sources"],
    )

    return result


@app.post("/report/dataset", response_model=ReportOut)
def report_dataset(
    topic: str = Query(...),
    document_ids: list[int] | None = Query(default=None),
    source: str = Query(default="auto"),
    integration_id: int | None = Query(default=None),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if document_ids:
        owned_count = (
            db.query(Document)
            .filter(Document.id.in_(document_ids), Document.user_id == current_user.id)
            .count()
        )
        if owned_count != len(document_ids):
            raise HTTPException(status_code=404, detail="One or more documents not found")

    settings = get_or_create_settings(db, current_user.id)
    plan = current_user.plan or "community"
    selected_integration = resolve_report_integration(db, current_user, plan, source, integration_id)

    if selected_integration and plan == "community":
        consume_community_credit(current_user, db)

    full_answer = []
    sources = []
    with RETRIEVAL_GATE(plan_priority(plan)):
        for event in generate_answer_stream(
            topic,
            user_id=current_user.id,
            document_ids=document_ids,
            ollama_url=settings.ollama_url,
            llm_model=settings.llm_model,
            external_provider=selected_integration.provider_name if selected_integration else None,
            external_api_key=selected_integration.api_key if selected_integration else None,
            external_base_url=selected_integration.base_url if selected_integration else None,
            external_model=selected_integration.model if selected_integration else None,
        ):
            if event["type"] == "sources":
                sources = event["sources"]
            elif event["type"] == "token":
                full_answer.append(event["content"])

    qa_pairs = [{"question": topic, "answer": "".join(full_answer), "sources": sources}]

    result = generate_report(
        qa_pairs,
        ollama_url=settings.ollama_url,
        llm_model=settings.llm_model,
        external_provider=selected_integration.provider_name if selected_integration else None,
        external_api_key=selected_integration.api_key if selected_integration else None,
        external_base_url=selected_integration.base_url if selected_integration else None,
        external_model=selected_integration.model if selected_integration else None,
    )

    save_generated_report(
        db,
        current_user,
        session_id=None,
        topic=topic,
        title=topic,
        report=result["report"],
        sources=result["sources"],
    )

    return result


def save_generated_report(db, current_user, *, session_id, title, report, sources, topic=None):
    """Upserts the saved report: one row per (user, session) for chat-based
    reports, one row per user with session_id NULL for the latest dataset report."""
    query = db.query(GeneratedReport).filter(GeneratedReport.user_id == current_user.id)
    query = (
        query.filter(GeneratedReport.session_id == session_id)
        if session_id is not None
        else query.filter(GeneratedReport.session_id.is_(None))
    )
    existing = query.first()

    if existing:
        existing.title = title
        existing.report = report
        existing.sources = ",".join(sources)
        existing.topic = topic
    else:
        existing = GeneratedReport(
            user_id=current_user.id,
            session_id=session_id,
            topic=topic,
            title=title,
            report=report,
            sources=",".join(sources),
        )
        db.add(existing)

    db.commit()
    return existing


@app.get("/reports", response_model=list[GeneratedReportOut])
def list_generated_reports(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    return (
        db.query(GeneratedReport)
        .filter(GeneratedReport.user_id == current_user.id)
        .order_by(GeneratedReport.updated_at.desc())
        .all()
    )


def parse_report_lines(report_text: str):
    """Splits the model's report text into (is_bullet, text) lines, stripping
    any leading markdown bullet marker the model added despite being told not to."""
    lines = []
    for raw_line in report_text.split("\n"):
        line = raw_line.strip()
        if not line:
            continue
        is_bullet = line.startswith(("* ", "- ", "• "))
        if is_bullet:
            line = line[2:].strip()
        lines.append((is_bullet, line))
    return lines


@app.post("/report/export")
def export_report(
    export_in: ReportExportIn,
    export_format: str = Query(..., alias="format"),
    current_user: User = Depends(get_current_user),
):
    if export_format not in ("docx", "pdf"):
        raise HTTPException(status_code=400, detail="Unsupported export format.")

    title = (export_in.title or "Report").strip() or "Report"
    lines = parse_report_lines(export_in.report)

    if export_format == "docx":
        from docx import Document as ExportDocxDocument

        doc = ExportDocxDocument()
        doc.add_heading(title, level=1)
        for is_bullet, text in lines:
            doc.add_paragraph(text, style="List Bullet" if is_bullet else None)

        if export_in.sources:
            doc.add_heading("Sources", level=2)
            for source_name in export_in.sources:
                doc.add_paragraph(source_name, style="List Bullet")

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=cognexa_report.docx"},
        )

    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    def escape(text):
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=LETTER)
    styles = getSampleStyleSheet()
    story = [Paragraph(escape(title), styles["Title"]), Spacer(1, 12)]

    for is_bullet, text in lines:
        prefix = "• " if is_bullet else ""
        story.append(Paragraph(f"{prefix}{escape(text)}", styles["BodyText"]))
        story.append(Spacer(1, 8))

    if export_in.sources:
        story.append(Spacer(1, 12))
        story.append(Paragraph("Sources", styles["Heading2"]))
        for source_name in export_in.sources:
            story.append(Paragraph(f"• {escape(source_name)}", styles["BodyText"]))

    doc.build(story)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=cognexa_report.pdf"},
    )


@app.get("/settings", response_model=SettingsOut)
def read_settings(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    return get_or_create_settings(db, current_user.id)


@app.put("/settings", response_model=SettingsOut)
def update_settings(
    settings_in: SettingsIn,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    settings = get_or_create_settings(db, current_user.id)

    for field, value in settings_in.model_dump().items():
        setattr(settings, field, value)

    db.commit()
    db.refresh(settings)

    return settings


def integration_to_out(integration: Integration) -> dict:
    return {
        "id": integration.id,
        "provider_name": integration.provider_name,
        "base_url": integration.base_url,
        "model": integration.model,
        "connected": bool(integration.api_key or integration.base_url),
        "created_at": integration.created_at,
    }


@app.get("/integrations", response_model=list[IntegrationOut])
def list_integrations(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    integrations = (
        db.query(Integration)
        .filter(Integration.user_id == current_user.id)
        .order_by(Integration.created_at.asc())
        .all()
    )
    return [integration_to_out(i) for i in integrations]


@app.post("/integrations", response_model=IntegrationOut)
def create_integration(
    integration_in: IntegrationIn,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    plan = current_user.plan or "community"

    if plan == "community" and not is_allowed_community_integration(
        integration_in.provider_name, integration_in.model, integration_in.base_url
    ):
        raise HTTPException(
            status_code=402,
            detail=(
                "The Community plan can only connect OpenRouter's free models. "
                "Upgrade your plan to use any provider and model."
            ),
        )

    existing_count = (
        db.query(Integration).filter(Integration.user_id == current_user.id).count()
    )
    max_integrations = INTEGRATION_LIMITS.get(plan, INTEGRATION_LIMITS["community"])
    if max_integrations is not None and existing_count >= max_integrations:
        raise HTTPException(
            status_code=402,
            detail=(
                f"{PLAN_DISPLAY_NAMES.get(plan, plan.capitalize())} plan allows up to "
                f"{max_integrations} saved integration(s). Remove one or upgrade your plan."
            ),
        )

    integration = Integration(
        user_id=current_user.id,
        provider_name=integration_in.provider_name,
        api_key=integration_in.api_key or None,
        base_url=integration_in.base_url or None,
        model=integration_in.model or None,
    )
    db.add(integration)
    db.commit()
    db.refresh(integration)

    return integration_to_out(integration)


@app.delete("/integrations/{integration_id}")
def delete_integration(
    integration_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    integration = (
        db.query(Integration)
        .filter(Integration.id == integration_id, Integration.user_id == current_user.id)
        .first()
    )
    if not integration:
        raise HTTPException(status_code=404, detail="Integration not found")

    db.delete(integration)
    db.commit()

    return {"deleted": integration_id}


def data_source_to_out(connection: DataSourceConnection) -> dict:
    return {
        "id": connection.id,
        "source_name": connection.source_name,
        "connected": bool(connection.credential),
        "status": connection.status,
        "status_message": connection.status_message,
        "last_synced_at": connection.last_synced_at,
        "synced_size_bytes": connection.synced_size_bytes or 0,
        # Never echo the credential (service account key) back to the client.
        "config": json.loads(connection.config) if connection.config else None,
        "created_at": connection.created_at,
    }


@app.get("/data-sources", response_model=list[DataSourceConnectionOut])
def list_data_sources(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    connections = (
        db.query(DataSourceConnection)
        .filter(DataSourceConnection.user_id == current_user.id)
        .order_by(DataSourceConnection.created_at.asc())
        .all()
    )
    return [data_source_to_out(c) for c in connections]


def _ensure_can_add_data_source(db: Session, current_user: User):
    plan = current_user.plan or "community"
    limits = PLAN_LIMITS.get(plan, PLAN_LIMITS["community"])
    max_apps = limits["max_apps"]

    existing_count = (
        db.query(DataSourceConnection)
        .filter(DataSourceConnection.user_id == current_user.id)
        .count()
    )
    if max_apps is not None and existing_count >= max_apps:
        raise HTTPException(
            status_code=402,
            detail=(
                f"{PLAN_DISPLAY_NAMES.get(plan, plan.capitalize())} plan allows up to "
                f"{max_apps} connected app(s). Remove one or upgrade your plan."
            ),
        )


GOOGLE_DRIVE_OAUTH_STATE_PURPOSE = "gdrive_oauth"
GOOGLE_DRIVE_OAUTH_STATE_MINUTES = 10


@app.get("/data-sources/google-drive/authorize", response_model=GoogleDriveAuthorizeOut)
def authorize_google_drive(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    _ensure_can_add_data_source(db, current_user)

    expire = datetime.now(timezone.utc) + timedelta(minutes=GOOGLE_DRIVE_OAUTH_STATE_MINUTES)
    state = jwt.encode(
        {"sub": str(current_user.id), "purpose": GOOGLE_DRIVE_OAUTH_STATE_PURPOSE, "exp": expire},
        SECRET_KEY,
        algorithm=ALGORITHM,
    )

    try:
        url = get_authorization_url(state)
    except RuntimeError as exc:
        raise HTTPException(status_code=500, detail=str(exc))

    return {"url": url}


@app.get("/data-sources/google-drive/callback")
def google_drive_callback(
    code: str | None = None,
    state: str | None = None,
    error: str | None = None,
    db: Session = Depends(get_db),
):
    frontend_url = os.getenv("FRONTEND_URL", "http://localhost:3000")
    redirect_target = f"{frontend_url}/settings/data-sources"

    if error:
        return RedirectResponse(f"{redirect_target}?error={error}")
    if not code or not state:
        return RedirectResponse(f"{redirect_target}?error=missing_code_or_state")

    try:
        payload = jwt.decode(state, SECRET_KEY, algorithms=[ALGORITHM])
        if payload.get("purpose") != GOOGLE_DRIVE_OAUTH_STATE_PURPOSE:
            raise JWTError("wrong purpose")
        user_id = int(payload["sub"])
    except (JWTError, KeyError, ValueError):
        return RedirectResponse(f"{redirect_target}?error=invalid_or_expired_state")

    current_user = db.query(User).filter(User.id == user_id).first()
    if not current_user:
        return RedirectResponse(f"{redirect_target}?error=user_not_found")

    try:
        _ensure_can_add_data_source(db, current_user)
    except HTTPException as exc:
        return RedirectResponse(f"{redirect_target}?error={exc.detail}")

    try:
        credentials = exchange_code_for_credentials(code)
        account_email = fetch_account_email(credentials)
        encrypted_tokens = encode_connection_tokens(credentials)
    except Exception as exc:
        return RedirectResponse(f"{redirect_target}?error=google_auth_failed")

    connection = DataSourceConnection(
        user_id=current_user.id,
        source_name="Google Drive",
        credential=encrypted_tokens,
        config=json.dumps({"sync_deleted": True, "account_email": account_email}),
        status="connected",
        status_message=f"Connected as {account_email}" if account_email else "Connected",
    )
    db.add(connection)
    db.commit()

    return RedirectResponse(f"{redirect_target}?connected=1")


@app.post("/data-sources", response_model=DataSourceConnectionOut)
def create_data_source(
    connection_in: DataSourceConnectionIn,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if connection_in.source_name == "Google Drive":
        raise HTTPException(
            status_code=400,
            detail="Google Drive connections are created via GET /data-sources/google-drive/authorize.",
        )

    _ensure_can_add_data_source(db, current_user)

    connection = DataSourceConnection(
        user_id=current_user.id,
        source_name=connection_in.source_name,
        credential=connection_in.credential or None,
        config=json.dumps(connection_in.config.model_dump()) if connection_in.config else None,
        status=None,
    )
    db.add(connection)
    db.commit()
    db.refresh(connection)

    return data_source_to_out(connection)


@app.delete("/data-sources/{connection_id}")
def delete_data_source(
    connection_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    connection = (
        db.query(DataSourceConnection)
        .filter(DataSourceConnection.id == connection_id, DataSourceConnection.user_id == current_user.id)
        .first()
    )
    if not connection:
        raise HTTPException(status_code=404, detail="Data source connection not found")

    # Documents already synced from this connection stay in the dataset --
    # only unlink them from the connection being removed (its FK would
    # otherwise block deletion).
    db.query(Document).filter(Document.data_source_id == connection.id).update(
        {Document.data_source_id: None}
    )

    db.delete(connection)
    db.commit()

    return {"deleted": connection_id}


@app.post("/data-sources/{connection_id}/sync", response_model=DataSourceSyncStartedOut)
def sync_data_source(
    connection_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    connection = (
        db.query(DataSourceConnection)
        .filter(DataSourceConnection.id == connection_id, DataSourceConnection.user_id == current_user.id)
        .first()
    )
    if not connection:
        raise HTTPException(status_code=404, detail="Data source connection not found")

    if connection.source_name != "Google Drive":
        raise HTTPException(status_code=400, detail="Syncing isn't supported for this data source yet.")

    if connection.status == "syncing":
        raise HTTPException(status_code=409, detail="A sync is already in progress for this connection.")

    plan = current_user.plan or "community"
    limits = PLAN_LIMITS.get(plan, PLAN_LIMITS["community"])
    document_count, storage_bytes = get_plan_usage(db, current_user.id)
    if limits["max_documents"] is not None and document_count >= limits["max_documents"]:
        raise HTTPException(
            status_code=402,
            detail=(
                f"{PLAN_DISPLAY_NAMES.get(plan, plan.capitalize())} plan limit reached "
                f"({limits['max_documents']} documents). Upgrade your plan to sync more."
            ),
        )

    settings = get_or_create_settings(db, current_user.id)

    connection.status = "syncing"
    connection.status_message = "Starting sync..."
    db.commit()

    INDEX_POOL.submit(
        plan_priority(plan),
        sync_google_drive_connection_job,
        connection.id,
        UPLOAD_FOLDER,
        settings.chunk_size,
        settings.chunk_overlap,
    )

    return {"started": True, "status": "syncing"}


def _get_google_drive_connection(connection_id: int, current_user: User, db: Session) -> DataSourceConnection:
    connection = (
        db.query(DataSourceConnection)
        .filter(DataSourceConnection.id == connection_id, DataSourceConnection.user_id == current_user.id)
        .first()
    )
    if not connection:
        raise HTTPException(status_code=404, detail="Data source connection not found")
    if connection.source_name != "Google Drive":
        raise HTTPException(status_code=400, detail="This action is only supported for Google Drive.")
    return connection


@app.get("/data-sources/{connection_id}/google-drive/picker-token", response_model=GoogleDrivePickerTokenOut)
def get_google_drive_picker_token(
    connection_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    connection = _get_google_drive_connection(connection_id, current_user, db)

    try:
        access_token, expires_in, refreshed_credential = get_picker_access_token(connection)
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"Failed to get a Google Drive access token: {exc}")

    if refreshed_credential:
        connection.credential = refreshed_credential
        db.commit()

    return {"access_token": access_token, "expires_in": expires_in}


@app.patch("/data-sources/{connection_id}/google-drive/folders", response_model=DataSourceConnectionOut)
def update_google_drive_folders(
    connection_id: int,
    folders_in: GoogleDriveFoldersIn,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    connection = _get_google_drive_connection(connection_id, current_user, db)

    config = json.loads(connection.config or "{}")
    config["folders"] = [f.model_dump() for f in folders_in.folders]
    connection.config = json.dumps(config)
    db.commit()
    db.refresh(connection)

    return data_source_to_out(connection)


def chat_channel_to_out(channel: ChatChannel) -> dict:
    return {
        "id": channel.id,
        "channel_name": channel.channel_name,
        "connected": bool(channel.bot_token),
        "bot_username": channel.bot_username,
        "created_at": channel.created_at,
    }


@app.get("/chat-channels", response_model=list[ChatChannelOut])
def list_chat_channels(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    channels = (
        db.query(ChatChannel)
        .filter(ChatChannel.user_id == current_user.id)
        .order_by(ChatChannel.created_at.asc())
        .all()
    )
    return [chat_channel_to_out(c) for c in channels]


@app.post("/chat-channels", response_model=ChatChannelOut)
def create_chat_channel(
    channel_in: ChatChannelIn,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    plan = current_user.plan or "community"
    limits = PLAN_LIMITS.get(plan, PLAN_LIMITS["community"])
    max_chat_channels = limits["max_chat_channels"]

    existing_count = (
        db.query(ChatChannel).filter(ChatChannel.user_id == current_user.id).count()
    )
    if max_chat_channels is not None and existing_count >= max_chat_channels:
        raise HTTPException(
            status_code=402,
            detail=(
                f"{PLAN_DISPLAY_NAMES.get(plan, plan.capitalize())} plan allows up to "
                f"{max_chat_channels} connected chat channel(s). Remove one or upgrade your plan."
            ),
        )

    channel = ChatChannel(
        user_id=current_user.id,
        channel_name=channel_in.channel_name,
    )

    if channel_in.channel_name.strip().lower() == "telegram" and channel_in.bot_token:
        raw_token = channel_in.bot_token.strip()
        try:
            bot_info = telegram_get_me(raw_token)
        except ValueError as exc:
            raise HTTPException(status_code=400, detail=str(exc))

        # request.base_url can't be trusted here -- behind Render's proxy it
        # reflects the app's own internal address (seen in practice as a
        # non-public port, and even as a loopback IP Telegram rejects
        # outright), not the bot's actual public URL. Require it explicitly
        # instead of guessing wrong in a way that only surfaces as a cryptic
        # Telegram API error.
        base_url = os.getenv("API_PUBLIC_URL")
        if not base_url:
            raise HTTPException(
                status_code=500,
                detail=(
                    "API_PUBLIC_URL is not set on the backend. Set it to this API's public "
                    "HTTPS URL (e.g. https://cognexa-api.onrender.com) in Render's environment "
                    "variables, then try connecting Telegram again."
                ),
            )

        parsed = urlsplit(base_url.rstrip("/"))
        is_local = parsed.hostname in ("127.0.0.1", "localhost", "::1")
        if not parsed.hostname or (parsed.scheme != "https" and not is_local):
            raise HTTPException(
                status_code=500,
                detail="API_PUBLIC_URL must be a full https:// URL, e.g. https://cognexa-api.onrender.com",
            )

        webhook_secret = generate_webhook_secret()

        if is_local:
            # Telegram's servers can never reach 127.0.0.1/localhost, no
            # matter what URL we hand them -- registering a real webhook
            # here would just fail (or silently go nowhere). Save the
            # channel locally without calling Telegram's API, so the rest of
            # the connect flow (token storage, DB rows, UI) is still
            # testable in local dev; use a public HTTPS tunnel (e.g. ngrok)
            # pointed at API_PUBLIC_URL for an actually-working bot locally.
            pass
        else:
            # Telegram only accepts a webhook on 443/80/88/8443 (or no
            # explicit port, i.e. the scheme's default) -- drop any other
            # port rather than send a URL Telegram will reject.
            netloc = parsed.hostname if parsed.port not in (None, 443, 80, 88, 8443) else parsed.netloc
            webhook_url = f"https://{netloc}/channels/telegram/webhook/{webhook_secret}"
            try:
                telegram_set_webhook(raw_token, webhook_url, webhook_secret)
            except ValueError as exc:
                raise HTTPException(status_code=400, detail=str(exc))

        channel.bot_token = encrypt_str(raw_token)
        channel.bot_username = bot_info.get("username")
        channel.webhook_secret = webhook_secret
    elif channel_in.bot_token:
        channel.bot_token = encrypt_str(channel_in.bot_token.strip())

    db.add(channel)
    db.commit()
    db.refresh(channel)

    return chat_channel_to_out(channel)


@app.get("/chat-channels/{channel_id}/telegram/test")
def test_telegram_channel(
    channel_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    channel = (
        db.query(ChatChannel)
        .filter(ChatChannel.id == channel_id, ChatChannel.user_id == current_user.id)
        .first()
    )
    if not channel:
        raise HTTPException(status_code=404, detail="Chat channel not found")
    if channel.channel_name.strip().lower() != "telegram" or not channel.bot_token:
        raise HTTPException(status_code=400, detail="This channel isn't a connected Telegram bot.")

    try:
        info = telegram_get_webhook_info(decrypt_str(channel.bot_token))
    except ValueError as exc:
        raise HTTPException(status_code=502, detail=str(exc))

    expected_url = (
        f"/channels/telegram/webhook/{channel.webhook_secret}" if channel.webhook_secret else None
    )
    webhook_registered = bool(info.get("url"))
    webhook_matches = bool(expected_url) and info.get("url", "").endswith(expected_url)

    if not webhook_registered:
        status_message = "No webhook is registered with Telegram for this bot."
    elif not webhook_matches:
        status_message = "A webhook is registered, but it doesn't point at this backend."
    elif info.get("last_error_message"):
        status_message = f"Telegram couldn't deliver the last update: {info['last_error_message']}"
    else:
        status_message = "Webhook is registered and healthy."

    return {
        "ok": webhook_registered and webhook_matches and not info.get("last_error_message"),
        "status_message": status_message,
        "webhook_url": info.get("url") or None,
        "pending_update_count": info.get("pending_update_count", 0),
        "last_error_message": info.get("last_error_message"),
        "last_error_date": info.get("last_error_date"),
    }


@app.delete("/chat-channels/{channel_id}")
def delete_chat_channel(
    channel_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    channel = (
        db.query(ChatChannel)
        .filter(ChatChannel.id == channel_id, ChatChannel.user_id == current_user.id)
        .first()
    )
    if not channel:
        raise HTTPException(status_code=404, detail="Chat channel not found")

    if channel.webhook_secret and channel.bot_token:
        try:
            telegram_delete_webhook(decrypt_str(channel.bot_token))
        except Exception:
            pass

    db.query(TelegramChatLink).filter(TelegramChatLink.channel_id == channel.id).delete()
    db.delete(channel)
    db.commit()

    return {"deleted": channel_id}


@app.post("/channels/telegram/webhook/{webhook_secret}")
async def telegram_webhook(
    webhook_secret: str,
    request: Request,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
):
    channel = (
        db.query(ChatChannel).filter(ChatChannel.webhook_secret == webhook_secret).first()
    )
    if not channel:
        # Always 200 so Telegram doesn't treat a stale/removed channel as a
        # delivery failure and keep retrying.
        return {"ok": True}

    try:
        update = await request.json()
    except Exception:
        return {"ok": True}

    channel_id = channel.id

    def run_update():
        # Use a dedicated session/row instead of the request-scoped `db` and
        # `channel` above -- those belong to the webhook request, which has
        # already returned its response by the time this runs.
        background_db = SessionLocal()
        try:
            logger.info("Processing Telegram update for channel %s", channel_id)
            channel_row = (
                background_db.query(ChatChannel).filter(ChatChannel.id == channel_id).first()
            )
            if channel_row:
                handle_telegram_update(background_db, channel_row, update, RETRIEVAL_GATE, plan_priority)
            logger.info("Finished Telegram update for channel %s", channel_id)
        except Exception:
            logger.exception("Telegram webhook handling failed for channel %s", channel_id)
        finally:
            background_db.close()

    # Acknowledge Telegram immediately instead of blocking on RAG generation
    # (which can take longer than Telegram/Render's proxy will wait, and on
    # its very first call also pays torch/sentence-transformers' load cost)
    # -- a slow or crashed synchronous response here is what surfaces to
    # Telegram as "Wrong response from the webhook: 502 Bad Gateway". The
    # actual reply still gets sent via sendMessage once this finishes, just
    # decoupled from the webhook's own response.
    background_tasks.add_task(run_update)

    return {"ok": True}


@app.get("/stats", response_model=StatsOut)
def get_stats(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    documents = db.query(Document).filter(Document.user_id == current_user.id).all()

    total_documents = len(documents)
    total_chunks = sum(doc.chunks or 0 for doc in documents)

    storage_bytes = sum(doc.size_bytes or 0 for doc in documents)

    today = datetime.now(timezone.utc).date()
    questions_today = (
        db.query(ChatMessage)
        .filter(ChatMessage.user_id == current_user.id)
        .all()
    )
    questions_today = sum(
        1
        for msg in questions_today
        if msg.created_at and msg.created_at.date() == today
    )

    return {
        "total_documents": total_documents,
        "total_chunks": total_chunks,
        "questions_today": questions_today,
        "storage_bytes": storage_bytes,
    }


@app.get("/data/export")
def export_knowledge_base(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    documents = (
        db.query(Document)
        .filter(Document.user_id == current_user.id)
        .order_by(Document.created_at.asc())
        .all()
    )

    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        manifest = []
        for document in documents:
            manifest.append({
                "id": document.id,
                "filename": document.filename,
                "file_type": document.file_type,
                "chunks": document.chunks,
                "size_bytes": document.size_bytes,
                "page_count": document.page_count,
                "created_at": document.created_at.isoformat() if document.created_at else None,
            })
            if document.file_path and os.path.exists(document.file_path):
                zf.write(document.file_path, arcname=f"documents/{document.filename}")

        zf.writestr("manifest.json", json.dumps(manifest, indent=2))

    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/zip",
        headers={"Content-Disposition": "attachment; filename=cognexa_knowledge_base.zip"},
    )


@app.get("/data/backup")
def backup_account(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    settings = get_or_create_settings(db, current_user.id)
    documents = db.query(Document).filter(Document.user_id == current_user.id).all()
    chat_messages = db.query(ChatMessage).filter(ChatMessage.user_id == current_user.id).all()
    integrations = db.query(Integration).filter(Integration.user_id == current_user.id).all()

    def iso(value):
        return value.isoformat() if value else None

    backup = {
        "backup_version": 1,
        "exported_at": datetime.now(timezone.utc).isoformat(),
        "user": {
            "name": current_user.name,
            "email": current_user.email,
            "plan": current_user.plan,
        },
        "settings": {
            "ollama_url": settings.ollama_url,
            "llm_model": settings.llm_model,
            "embedding_model": settings.embedding_model,
            "chunk_size": settings.chunk_size,
            "chunk_overlap": settings.chunk_overlap,
            "theme": settings.theme,
            "email_notifications": settings.email_notifications,
            "notify_indexing_finished": settings.notify_indexing_finished,
            "notify_failed_uploads": settings.notify_failed_uploads,
            "notify_ai_errors": settings.notify_ai_errors,
            "notify_weekly_report": settings.notify_weekly_report,
        },
        # Document metadata only -- file content is not included here, use
        # "Export Knowledge Base" for the actual files.
        "documents": [
            {
                "filename": d.filename,
                "file_type": d.file_type,
                "chunks": d.chunks,
                "preview": d.preview,
                "size_bytes": d.size_bytes,
                "page_count": d.page_count,
                "created_at": iso(d.created_at),
            }
            for d in documents
        ],
        "chat_messages": [
            {
                "question": m.question,
                "answer": m.answer,
                "sources": m.sources,
                "created_at": iso(m.created_at),
            }
            for m in chat_messages
        ],
        # API keys are never included in a backup for security -- restoring an
        # integration re-creates the provider/model entry but you'll need to
        # re-enter its key afterward.
        "integrations": [
            {
                "provider_name": i.provider_name,
                "base_url": i.base_url,
                "model": i.model,
            }
            for i in integrations
        ],
    }

    return Response(
        content=json.dumps(backup, indent=2),
        media_type="application/json",
        headers={"Content-Disposition": "attachment; filename=cognexa_backup.json"},
    )


@app.post("/data/restore")
async def restore_account(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    try:
        backup = json.loads(await file.read())
    except (json.JSONDecodeError, UnicodeDecodeError):
        raise HTTPException(status_code=400, detail="That file isn't a valid backup export.")

    if not isinstance(backup, dict) or "settings" not in backup:
        raise HTTPException(status_code=400, detail="That file isn't a valid Cognexa backup.")

    # Restoring replaces current settings, chat history, and integrations with
    # the backup's. Documents are not restored here -- the backup only stores
    # document metadata, not file content, so there's nothing to re-index.
    settings = get_or_create_settings(db, current_user.id)
    for field, value in (backup.get("settings") or {}).items():
        if hasattr(settings, field):
            setattr(settings, field, value)

    db.query(ChatMessage).filter(ChatMessage.user_id == current_user.id).delete()
    db.query(ChatSession).filter(ChatSession.user_id == current_user.id).delete()

    restored_messages = backup.get("chat_messages") or []
    if restored_messages:
        restored_session = ChatSession(user_id=current_user.id, title="Restored conversation")
        db.add(restored_session)
        db.flush()

        for m in restored_messages:
            db.add(ChatMessage(
                user_id=current_user.id,
                session_id=restored_session.id,
                question=m.get("question", ""),
                answer=m.get("answer", ""),
                sources=m.get("sources"),
            ))

    plan = current_user.plan or "community"
    max_integrations = INTEGRATION_LIMITS.get(plan, INTEGRATION_LIMITS["community"])
    db.query(Integration).filter(Integration.user_id == current_user.id).delete()
    restored_integrations = 0
    for i in backup.get("integrations") or []:
        if max_integrations is not None and restored_integrations >= max_integrations:
            break
        if plan == "community" and not is_allowed_community_integration(
            i.get("provider_name"), i.get("model"), i.get("base_url")
        ):
            continue
        db.add(Integration(
            user_id=current_user.id,
            provider_name=i.get("provider_name", "Cline"),
            base_url=i.get("base_url"),
            model=i.get("model"),
        ))
        restored_integrations += 1

    db.commit()

    return {
        "restored_chat_messages": len(backup.get("chat_messages") or []),
        "restored_integrations": restored_integrations,
    }


@app.delete("/account")
def delete_account(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    documents = db.query(Document).filter(Document.user_id == current_user.id).all()
    for document in documents:
        if document.file_path and os.path.exists(document.file_path):
            os.remove(document.file_path)
        db.delete(document)

    delete_user_collection(current_user.id)

    db.query(GeneratedReport).filter(GeneratedReport.user_id == current_user.id).delete()
    db.query(ChatMessage).filter(ChatMessage.user_id == current_user.id).delete()
    db.query(ChatSession).filter(ChatSession.user_id == current_user.id).delete()
    db.query(Integration).filter(Integration.user_id == current_user.id).delete()
    db.query(Settings).filter(Settings.user_id == current_user.id).delete()
    db.query(DataSourceConnection).filter(DataSourceConnection.user_id == current_user.id).delete()
    db.query(ChatChannel).filter(ChatChannel.user_id == current_user.id).delete()
    db.execute(text("DELETE FROM activity_log WHERE user_id = :user_id"), {"user_id": current_user.id})

    db.delete(current_user)
    db.commit()

    return {"deleted": True}


if __name__ == "__main__":
    import uvicorn

    port = int(os.getenv("PORT", 10000))
    uvicorn.run("app.main:app", host="0.0.0.0", port=port)
