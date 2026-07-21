import json
import os
from datetime import datetime, timezone
import io
import zipfile
from typing import Any

from fastapi import APIRouter, Depends, HTTPException, Query, UploadFile, File, Form
from sqlalchemy import desc, func
from sqlalchemy.orm import Session

from app.auth import get_current_user
from app.database import get_db
from app.models import (
    SDLCActivityLog,
    Pipeline,
    PipelineExecution,
    PipelineStep,
    Project,
    ProjectArtifact,
    SDLCStage,
    ToolIntegration,
    User,
    Integration,
    Settings,
    Document,
)
from app.schemas import (
    ArtifactGenerateIn,
    PipelineExecutionOut,
    PipelineIn,
    PipelineOut,
    PipelineStepIn,
    ProjectArtifactIn,
    ProjectArtifactOut,
    ProjectArtifactContentOut,
    ProjectArtifactUpdateIn,
    ProjectIn,
    ProjectOut,
    SDLCStageIn,
    SDLCStageOut,
    SDLCActivityLogOut,
    ToolIntegrationIn,
    ToolIntegrationOut,
)
from app.crypto import encrypt_str, decrypt_str
from app.sdlc_ai import generate_artifact, generate_sdlc_plan
from app.sdlc_generator import start_async_generation

router = APIRouter(prefix="/sdlc", tags=["sdlc"])


# ============================================================
# Activity Log Helper
# ============================================================


def log_activity(
    db: Session,
    user_id: int,
    action: str,
    project_id: int | None = None,
    stage_id: int | None = None,
    entity_type: str | None = None,
    entity_id: int | None = None,
    description: str | None = None,
    metadata: dict | None = None,
):
    log = SDLCActivityLog(
        user_id=user_id,
        project_id=project_id,
        stage_id=stage_id,
        action=action,
        entity_type=entity_type,
        entity_id=entity_id,
        description=description,
        metadata=metadata,
    )
    db.add(log)
    db.commit()


def get_project_or_404(project_id: int, user_id: int, db: Session):
    """Fetch a project scoped to the current user, or raise 404."""
    project = (
        db.query(Project)
        .filter(Project.id == project_id, Project.user_id == user_id)
        .first()
    )
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    return project


def get_settings_or_default(db: Session, user_id: int):
    """Get user settings with defaults, creating them if they don't exist."""
    settings = db.query(Settings).filter(Settings.user_id == user_id).first()
    if not settings:
        settings = Settings(user_id=user_id)
        db.add(settings)
        db.commit()
        db.refresh(settings)
    return settings


# ============================================================
# Resolve AI provider from an Integration ID
# ============================================================


def resolve_artifact_provider(
    db: Session, user_id: int, integration_id: int | None
):
    """Returns (external_provider, external_api_key, external_base_url,
    external_model) from the given integration, or (None, None, None, None)
    if no integration is specified (falling back to Ollama)."""
    if integration_id is None:
        return None, None, None, None

    integration = (
        db.query(Integration)
        .filter(Integration.id == integration_id, Integration.user_id == user_id)
        .first()
    )
    if not integration:
        raise HTTPException(status_code=404, detail="Integration not found")
    if not integration.api_key:
        return None, None, None, None
    return (
        integration.provider_name,
        integration.api_key,
        integration.base_url,
        integration.model,
    )


# ============================================================
# Plan limits helper
# ============================================================


def get_plan_usage(db: Session, user_id: int):
    documents = db.query(Document).filter(Document.user_id == user_id).all()
    document_count = len(documents)
    storage_bytes = sum(doc.size_bytes or 0 for doc in documents)
    return document_count, storage_bytes


PLAN_LIMITS = {
    "community": {
        "max_documents": 25,
        "max_storage_bytes": 15 * 1024 * 1024,
        "max_apps": 2,
        "max_chat_channels": 1,
    },
    "pro": {
        "max_documents": 100,
        "max_storage_bytes": 10 * 1024 * 1024 * 1024,
        "max_apps": 10,
        "max_chat_channels": 5,
    },
    "team": {
        "max_documents": None,
        "max_storage_bytes": None,
        "max_apps": None,
        "max_chat_channels": None,
    },
}

PLAN_DISPLAY_NAMES = {
    "community": "Free",
    "pro": "Pro",
    "team": "Unlimited",
}


# ============================================================
# Projects
# ============================================================


@router.get("/projects", response_model=list[ProjectOut])
def list_projects(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    projects = (
        db.query(Project)
        .filter(Project.user_id == current_user.id)
        .order_by(Project.updated_at.desc())
        .all()
    )
    return projects


@router.post("/projects", response_model=ProjectOut)
def create_project(
    project_in: ProjectIn,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    project = Project(
        user_id=current_user.id,
        name=project_in.name,
        description=project_in.description,
        status=project_in.status,
        sdlc_stage=project_in.sdlc_stage,
        repository_url=project_in.repository_url,
        start_date=project_in.start_date,
        target_date=project_in.target_date,
    )
    db.add(project)
    db.commit()
    db.refresh(project)

    log_activity(
        db,
        user_id=current_user.id,
        project_id=project.id,
        action="project_created",
        entity_type="project",
        entity_id=project.id,
        description=f"Project '{project.name}' created",
    )

    return project


@router.post("/projects/generate", response_model=ProjectOut)
async def generate_project(
    name: str = Form(...),
    description: str = Form(""),
    integration_id: int | None = Form(None),
    documents: list[UploadFile] | None = File(default=None),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Create a project (status=generating) and kick off background AI generation.

    Returns immediately so the UI can redirect to the dashboard. The frontend
    polls GET /sdlc/projects/{id} to track progress via generation_stage/status.
    """
    plan = current_user.plan or "community"

    document_count, storage_bytes = get_plan_usage(db, current_user.id)
    limits = PLAN_LIMITS.get(plan, PLAN_LIMITS["community"])
    plan_display_name = PLAN_DISPLAY_NAMES.get(plan, plan.capitalize())

    if limits["max_documents"] is not None and document_count >= limits["max_documents"]:
        raise HTTPException(
            status_code=402,
            detail=f"{plan_display_name} plan limit reached ({limits['max_documents']} documents). Upgrade your plan for more.",
        )

    project = Project(
        user_id=current_user.id,
        name=name.strip(),
        description=description.strip() or None,
        status="generating",
        sdlc_stage="planning",
    )
    db.add(project)
    db.commit()
    db.refresh(project)

    document_texts: list[str] = []
    for upload in documents or []:
        if not upload.filename:
            continue
        contents = await upload.read()
        upload_folder = "app/uploads"
        os.makedirs(upload_folder, exist_ok=True)
        file_path = os.path.join(upload_folder, upload.filename)
        with open(file_path, "wb") as buffer:
            buffer.write(contents)

        from app.main import extract_document_text
        text, file_type, page_count = extract_document_text(file_path, upload.filename)
        if text and text.strip():
            document_texts.append(text)

    ext_provider, ext_key, ext_url, ext_model = resolve_artifact_provider(
        db, current_user.id, integration_id
    )
    settings = get_settings_or_default(db, current_user.id)

    log_activity(
        db,
        user_id=current_user.id,
        project_id=project.id,
        action="project_generation_started",
        entity_type="project",
        entity_id=project.id,
        description=f"Project '{project.name}' generation started (async)",
    )

    start_async_generation(
        project_id=project.id,
        user_id=current_user.id,
        name=name.strip(),
        description=description.strip(),
        document_texts=document_texts,
        ext_provider=ext_provider,
        ext_key=ext_key,
        ext_url=ext_url,
        ext_model=ext_model,
        ollama_url=settings.ollama_url,
        llm_model=settings.llm_model,
    )

    db.refresh(project)
    return project


@router.post("/projects/{project_id}/retry", response_model=ProjectOut)
def retry_generation(
    project_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Retry a failed generation for an existing project."""
    project = get_project_or_404(project_id, current_user.id, db)

    if project.status != "error":
        raise HTTPException(status_code=400, detail="Only failed projects can be retried.")

    project.status = "generating"
    project.generation_stage = "analyzing"
    project.generation_error = None
    project.generation_progress = None
    db.commit()

    settings = get_settings_or_default(db, current_user.id)

    log_activity(
        db,
        user_id=current_user.id,
        project_id=project.id,
        action="project_generation_retried",
        entity_type="project",
        entity_id=project.id,
        description=f"Generation retried for project '{project.name}'",
    )

    start_async_generation(
        project_id=project.id,
        user_id=current_user.id,
        name=project.name,
        description=project.description or "",
        document_texts=[],
        ext_provider=None,
        ext_key=None,
        ext_url=None,
        ext_model=None,
        ollama_url=settings.ollama_url,
        llm_model=settings.llm_model,
    )

    db.refresh(project)
    return project


@router.get("/projects/{project_id}", response_model=ProjectOut)
def get_project(
    project_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    return get_project_or_404(project_id, current_user.id, db)


@router.patch("/projects/{project_id}", response_model=ProjectOut)
def update_project(
    project_id: int,
    project_in: ProjectIn,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    project = get_project_or_404(project_id, current_user.id, db)

    project.name = project_in.name
    project.description = project_in.description
    project.status = project_in.status
    project.sdlc_stage = project_in.sdlc_stage
    project.repository_url = project_in.repository_url
    project.start_date = project_in.start_date
    project.target_date = project_in.target_date

    db.commit()
    db.refresh(project)

    log_activity(
        db,
        user_id=current_user.id,
        project_id=project.id,
        action="project_updated",
        entity_type="project",
        entity_id=project.id,
        description=f"Project '{project.name}' updated",
    )

    return project


@router.delete("/projects/{project_id}")
def delete_project(
    project_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    project = get_project_or_404(project_id, current_user.id, db)
    db.query(SDLCActivityLog).filter(SDLCActivityLog.project_id == project_id).delete()
    db.delete(project)
    db.commit()
    return {"deleted": project_id}


# ============================================================
# SDLC Stages
# ============================================================


@router.get("/projects/{project_id}/stages", response_model=list[SDLCStageOut])
def list_stages(
    project_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    get_project_or_404(project_id, current_user.id, db)
    stages = (
        db.query(SDLCStage)
        .filter(SDLCStage.project_id == project_id)
        .order_by(SDLCStage.priority.desc(), SDLCStage.created_at.asc())
        .all()
    )
    return stages


@router.post("/projects/{project_id}/stages", response_model=SDLCStageOut)
def create_stage(
    project_id: int,
    stage_in: SDLCStageIn,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    project = get_project_or_404(project_id, current_user.id, db)

    stage = SDLCStage(
        project_id=project_id,
        stage_type=stage_in.stage_type,
        name=stage_in.name,
        description=stage_in.description,
        priority=stage_in.priority,
        assigned_to=stage_in.assigned_to,
    )
    db.add(stage)
    db.commit()
    db.refresh(stage)

    log_activity(
        db,
        user_id=current_user.id,
        project_id=project_id,
        stage_id=stage.id,
        action="stage_created",
        entity_type="sdlc_stage",
        entity_id=stage.id,
        description=f"Stage '{stage.name}' created in project '{project.name}'",
    )

    return stage


@router.patch("/stages/{stage_id}", response_model=SDLCStageOut)
def update_stage(
    stage_id: int,
    stage_in: SDLCStageIn,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    stage = db.query(SDLCStage).filter(SDLCStage.id == stage_id).first()
    if not stage:
        raise HTTPException(status_code=404, detail="Stage not found")

    get_project_or_404(stage.project_id, current_user.id, db)

    stage.stage_type = stage_in.stage_type
    stage.name = stage_in.name
    stage.description = stage_in.description
    stage.priority = stage_in.priority
    stage.assigned_to = stage_in.assigned_to

    db.commit()
    db.refresh(stage)

    log_activity(
        db,
        user_id=current_user.id,
        project_id=stage.project_id,
        stage_id=stage.id,
        action="stage_updated",
        entity_type="sdlc_stage",
        entity_id=stage.id,
        description=f"Stage '{stage.name}' updated",
    )

    return stage


@router.patch("/stages/{stage_id}/status")
def update_stage_status(
    stage_id: int,
    status: str = Query(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    stage = db.query(SDLCStage).filter(SDLCStage.id == stage_id).first()
    if not stage:
        raise HTTPException(status_code=404, detail="Stage not found")

    get_project_or_404(stage.project_id, current_user.id, db)

    old_status = stage.status
    stage.status = status
    if status == "in_progress" and not stage.started_at:
        stage.started_at = datetime.now(timezone.utc)
    if status == "completed" and not stage.completed_at:
        stage.completed_at = datetime.now(timezone.utc)

    db.commit()
    db.refresh(stage)

    log_activity(
        db,
        user_id=current_user.id,
        project_id=stage.project_id,
        stage_id=stage.id,
        action="stage_status_changed",
        entity_type="sdlc_stage",
        entity_id=stage.id,
        description=f"Stage '{stage.name}' status changed to '{status}'",
        metadata={"old_status": old_status, "new_status": status},
    )

    return {"status": stage.status}


@router.delete("/stages/{stage_id}")
def delete_stage(
    stage_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    stage = db.query(SDLCStage).filter(SDLCStage.id == stage_id).first()
    if not stage:
        raise HTTPException(status_code=404, detail="Stage not found")

    get_project_or_404(stage.project_id, current_user.id, db)
    db.query(SDLCActivityLog).filter(SDLCActivityLog.stage_id == stage_id).delete()
    db.delete(stage)
    db.commit()

    return {"deleted": stage_id}


# ============================================================
# Project Artifacts
# ============================================================


@router.get("/projects/{project_id}/artifacts", response_model=list[ProjectArtifactOut])
def list_artifacts(
    project_id: int,
    stage_id: int | None = Query(None),
    artifact_type: str | None = Query(None),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    get_project_or_404(project_id, current_user.id, db)

    query = db.query(ProjectArtifact).filter(ProjectArtifact.project_id == project_id)
    if stage_id:
        query = query.filter(ProjectArtifact.stage_id == stage_id)
    if artifact_type:
        query = query.filter(ProjectArtifact.artifact_type == artifact_type)

    artifacts = query.order_by(ProjectArtifact.created_at.desc()).all()
    return artifacts


@router.post("/projects/{project_id}/artifacts", response_model=ProjectArtifactOut)
def create_artifact(
    project_id: int,
    artifact_in: ProjectArtifactIn,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    project = get_project_or_404(project_id, current_user.id, db)

    artifact = ProjectArtifact(
        project_id=project_id,
        stage_id=artifact_in.stage_id,
        artifact_type=artifact_in.artifact_type,
        name=artifact_in.name,
        description=artifact_in.description,
        artifact_metadata=artifact_in.metadata,
    )
    db.add(artifact)
    db.commit()
    db.refresh(artifact)

    log_activity(
        db,
        user_id=current_user.id,
        project_id=project_id,
        stage_id=artifact_in.stage_id,
        action="artifact_created",
        entity_type="project_artifact",
        entity_id=artifact.id,
        description=f"Artifact '{artifact.name}' created in project '{project.name}'",
    )

    return artifact


@router.get("/artifacts/{artifact_id}", response_model=ProjectArtifactOut)
def get_artifact(
    artifact_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    artifact = db.query(ProjectArtifact).filter(ProjectArtifact.id == artifact_id).first()
    if not artifact:
        raise HTTPException(status_code=404, detail="Artifact not found")

    get_project_or_404(artifact.project_id, current_user.id, db)
    return artifact


@router.patch("/artifacts/{artifact_id}", response_model=ProjectArtifactOut)
def update_artifact(
    artifact_id: int,
    artifact_in: ProjectArtifactUpdateIn,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    artifact = db.query(ProjectArtifact).filter(ProjectArtifact.id == artifact_id).first()
    if not artifact:
        raise HTTPException(status_code=404, detail="Artifact not found")

    get_project_or_404(artifact.project_id, current_user.id, db)

    if artifact_in.name is not None:
        artifact.name = artifact_in.name
    if artifact_in.description is not None:
        artifact.description = artifact_in.description
    if artifact_in.content is not None:
        artifact.content = artifact_in.content
    if artifact_in.metadata is not None:
        artifact.artifact_metadata = artifact_in.metadata

    artifact.version += 1

    db.commit()
    db.refresh(artifact)

    return artifact


@router.get("/artifacts/{artifact_id}/content", response_model=ProjectArtifactContentOut)
def get_artifact_content(
    artifact_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    artifact = db.query(ProjectArtifact).filter(ProjectArtifact.id == artifact_id).first()
    if not artifact:
        raise HTTPException(status_code=404, detail="Artifact not found")

    get_project_or_404(artifact.project_id, current_user.id, db)

    return ProjectArtifactContentOut(
        content=artifact.content,
        generated_by=artifact.generated_by,
    )


@router.delete("/artifacts/{artifact_id}")
def delete_artifact(
    artifact_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    artifact = db.query(ProjectArtifact).filter(ProjectArtifact.id == artifact_id).first()
    if not artifact:
        raise HTTPException(status_code=404, detail="Artifact not found")

    get_project_or_404(artifact.project_id, current_user.id, db)
    db.delete(artifact)
    db.commit()

    return {"deleted": artifact_id}


# ============================================================
# AI Artifact Generation
# ============================================================


@router.post(
    "/projects/{project_id}/generate",
    response_model=ProjectArtifactOut,
)
def generate_project_artifact(
    project_id: int,
    gen_in: ArtifactGenerateIn,
    integration_id: int | None = Query(default=None),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Generate an AI-powered artifact for the given project."""
    project = get_project_or_404(project_id, current_user.id, db)

    ext_provider, ext_key, ext_url, ext_model = resolve_artifact_provider(
        db, current_user.id, integration_id
    )
    settings = get_settings_or_default(db, current_user.id)

    result = generate_artifact(
        artifact_type=gen_in.artifact_type,
        name=gen_in.name,
        description=gen_in.description,
        artifact_prompt=gen_in.artifact_prompt,
        external_provider=ext_provider,
        external_api_key=ext_key,
        external_base_url=ext_url,
        external_model=ext_model,
        ollama_url=settings.ollama_url,
        llm_model=settings.llm_model,
    )

    artifact = ProjectArtifact(
        project_id=project_id,
        stage_id=gen_in.stage_id,
        artifact_type=gen_in.artifact_type,
        name=gen_in.name,
        description=gen_in.description,
        content=result["content"],
        generated_by=result["generated_by"],
        artifact_metadata=gen_in.metadata,
        file_type="text",
        size_bytes=len(result["content"].encode("utf-8")),
    )
    db.add(artifact)
    db.commit()
    db.refresh(artifact)

    log_activity(
        db,
        user_id=current_user.id,
        project_id=project_id,
        stage_id=gen_in.stage_id,
        action="artifact_generated",
        entity_type="project_artifact",
        entity_id=artifact.id,
        description=f"AI-generated '{artifact.name}' ({artifact.artifact_type}) using {result['generated_by']}",
        metadata={
            "generated_by": result["generated_by"],
            "artifact_type": gen_in.artifact_type,
        },
    )

    return artifact


# ============================================================
# Pipelines
# ============================================================


@router.get("/pipelines", response_model=list[PipelineOut])
def list_pipelines(
    project_id: int | None = Query(None),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    query = db.query(Pipeline).filter(Pipeline.user_id == current_user.id)
    if project_id:
        query = query.filter(Pipeline.project_id == project_id)
    pipelines = query.order_by(Pipeline.updated_at.desc()).all()
    return pipelines


@router.post("/pipelines", response_model=PipelineOut)
def create_pipeline(
    pipeline_in: PipelineIn,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if pipeline_in.project_id:
        get_project_or_404(pipeline_in.project_id, current_user.id, db)

    pipeline = Pipeline(
        user_id=current_user.id,
        project_id=pipeline_in.project_id,
        name=pipeline_in.name,
        description=pipeline_in.description,
        pipeline_type=pipeline_in.pipeline_type,
        config=pipeline_in.config,
        is_template=pipeline_in.is_template,
        is_active=pipeline_in.is_active,
    )
    db.add(pipeline)
    db.commit()
    db.refresh(pipeline)

    for step_in in pipeline_in.steps:
        step = PipelineStep(
            pipeline_id=pipeline.id,
            step_order=step_in.step_order,
            step_type=step_in.step_type,
            name=step_in.name,
            config=step_in.config,
            input_mapping=step_in.input_mapping,
            output_mapping=step_in.output_mapping,
            timeout_seconds=step_in.timeout_seconds,
            retry_count=step_in.retry_count,
        )
        db.add(step)

    db.commit()
    db.refresh(pipeline)

    log_activity(
        db,
        user_id=current_user.id,
        project_id=pipeline_in.project_id,
        action="pipeline_created",
        entity_type="pipeline",
        entity_id=pipeline.id,
        description=f"Pipeline '{pipeline.name}' created",
    )

    return pipeline


@router.get("/pipelines/{pipeline_id}", response_model=PipelineOut)
def get_pipeline(
    pipeline_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    pipeline = (
        db.query(Pipeline)
        .filter(Pipeline.id == pipeline_id, Pipeline.user_id == current_user.id)
        .first()
    )
    if not pipeline:
        raise HTTPException(status_code=404, detail="Pipeline not found")
    return pipeline


@router.post("/pipelines/{pipeline_id}/execute", response_model=PipelineExecutionOut)
def execute_pipeline(
    pipeline_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    pipeline = (
        db.query(Pipeline)
        .filter(Pipeline.id == pipeline_id, Pipeline.user_id == current_user.id)
        .first()
    )
    if not pipeline:
        raise HTTPException(status_code=404, detail="Pipeline not found")

    execution = PipelineExecution(
        pipeline_id=pipeline.id,
        user_id=current_user.id,
        project_id=pipeline.project_id,
        status="pending",
        triggered_by="manual",
    )
    db.add(execution)
    db.commit()
    db.refresh(execution)

    log_activity(
        db,
        user_id=current_user.id,
        project_id=pipeline.project_id,
        action="pipeline_executed",
        entity_type="pipeline_execution",
        entity_id=execution.id,
        description=f"Pipeline '{pipeline.name}' execution started",
        metadata={"pipeline_id": pipeline.id, "execution_id": execution.id},
    )

    return execution


@router.get("/pipelines/{pipeline_id}/executions", response_model=list[PipelineExecutionOut])
def list_executions(
    pipeline_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    pipeline = (
        db.query(Pipeline)
        .filter(Pipeline.id == pipeline_id, Pipeline.user_id == current_user.id)
        .first()
    )
    if not pipeline:
        raise HTTPException(status_code=404, detail="Pipeline not found")

    executions = (
        db.query(PipelineExecution)
        .filter(PipelineExecution.pipeline_id == pipeline_id)
        .order_by(PipelineExecution.created_at.desc())
        .all()
    )
    return executions


# ============================================================
# Tool Integrations
# ============================================================


@router.get("/tools", response_model=list[ToolIntegrationOut])
def list_tools(
    project_id: int | None = Query(None),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    query = db.query(ToolIntegration).filter(ToolIntegration.user_id == current_user.id)
    if project_id:
        query = query.filter(ToolIntegration.project_id == project_id)
    tools = query.order_by(ToolIntegration.created_at.desc()).all()
    return tools


@router.post("/tools", response_model=ToolIntegrationOut)
def create_tool(
    tool_in: ToolIntegrationIn,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if tool_in.project_id:
        get_project_or_404(tool_in.project_id, current_user.id, db)

    encrypted_creds = encrypt_str(tool_in.credentials_encrypted) if tool_in.credentials_encrypted else None

    tool = ToolIntegration(
        user_id=current_user.id,
        project_id=tool_in.project_id,
        tool_name=tool_in.tool_name,
        display_name=tool_in.display_name,
        config=tool_in.config,
        credentials_encrypted=encrypted_creds,
    )
    db.add(tool)
    db.commit()
    db.refresh(tool)

    log_activity(
        db,
        user_id=current_user.id,
        project_id=tool_in.project_id,
        action="tool_connected",
        entity_type="tool_integration",
        entity_id=tool.id,
        description=f"Tool '{tool.tool_name}' connected",
    )

    return tool


@router.patch("/tools/{tool_id}", response_model=ToolIntegrationOut)
def update_tool(
    tool_id: int,
    tool_in: ToolIntegrationIn,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    tool = (
        db.query(ToolIntegration)
        .filter(ToolIntegration.id == tool_id, ToolIntegration.user_id == current_user.id)
        .first()
    )
    if not tool:
        raise HTTPException(status_code=404, detail="Tool not found")

    tool.display_name = tool_in.display_name
    tool.config = tool_in.config
    if tool_in.credentials_encrypted:
        tool.credentials_encrypted = encrypt_str(tool_in.credentials_encrypted)

    db.commit()
    db.refresh(tool)

    return tool


@router.delete("/tools/{tool_id}")
def delete_tool(
    tool_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    tool = (
        db.query(ToolIntegration)
        .filter(ToolIntegration.id == tool_id, ToolIntegration.user_id == current_user.id)
        .first()
    )
    if not tool:
        raise HTTPException(status_code=404, detail="Tool not found")

    db.delete(tool)
    db.commit()

    return {"deleted": tool_id}


# ============================================================
# Activity Log
# ============================================================


@router.get("/activity", response_model=list[SDLCActivityLogOut])
def list_activity(
    project_id: int | None = Query(None),
    limit: int = Query(50, le=200),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    query = db.query(SDLCActivityLog).filter(SDLCActivityLog.user_id == current_user.id)
    if project_id:
        query = query.filter(SDLCActivityLog.project_id == project_id)
    logs = query.order_by(SDLCActivityLog.created_at.desc()).limit(limit).all()
    return logs


# ============================================================
# Artifact Downloads
# ============================================================

from fastapi.responses import Response, StreamingResponse
import re


def build_playwright_zip_bytes(files: dict) -> bytes:
    """Package a Playwright project file map ({path: content}) into a zip archive."""
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for path, file_content in files.items():
            zf.writestr(path, file_content or "")
    buffer.seek(0)
    return buffer.read()


def build_docx_bytes(title: str, content: str) -> bytes:
    """Render artifact/report text content into a .docx document."""
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_heading(title or "Untitled", level=1)
    for raw_line in (content or "").split("\n"):
        line = raw_line.strip()
        if not line:
            continue
        heading_match = re.match(r"^(#{1,3})\s+(.*)", line)
        if heading_match:
            doc.add_heading(heading_match.group(2), level=min(len(heading_match.group(1)) + 1, 4))
            continue
        is_bullet = line.startswith(("* ", "- ", "• "))
        if is_bullet:
            line = line[2:].strip()
        doc.add_paragraph(line, style="List Bullet" if is_bullet else None)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


@router.get("/artifacts/{artifact_id}/download")
def download_artifact_file(
    artifact_id: int,
    format: str = Query("docx"),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    artifact = db.query(ProjectArtifact).filter(ProjectArtifact.id == artifact_id).first()
    if not artifact:
        raise HTTPException(status_code=404, detail="Artifact not found")
    get_project_or_404(artifact.project_id, current_user.id, db)

    content = artifact.content or ""
    safe_name = re.sub(r"[^a-zA-Z0-9_-]", "_", artifact.name).strip("_") or "artifact"

    if artifact.file_type == "playwright_project":
        files = (artifact.artifact_metadata or {}).get("files") or {}
        return Response(
            content=build_playwright_zip_bytes(files),
            media_type="application/zip",
            headers={"Content-Disposition": f"attachment; filename={safe_name}.zip"},
        )

    if format == "json":
        payload = json.dumps({
            "id": artifact.id,
            "name": artifact.name,
            "artifact_type": artifact.artifact_type,
            "description": artifact.description,
            "content": content,
            "generated_by": artifact.generated_by,
            "version": artifact.version,
            "created_at": artifact.created_at.isoformat() if artifact.created_at else None,
        }, indent=2)
        media_type = "application/json"
        filename = f"{safe_name}.json"
        body = payload.encode("utf-8")
    elif format in ("md", "txt"):
        ext = "md" if format == "md" else "txt"
        media_type = "text/plain"
        filename = f"{safe_name}.{ext}"
        body = content.encode("utf-8")
    else:
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = f"{safe_name}.docx"
        body = build_docx_bytes(artifact.name, content)

    return Response(
        content=body,
        media_type=media_type,
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@router.get("/projects/{project_id}/artifacts/download-all")
def download_all_artifacts(
    project_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    project = get_project_or_404(project_id, current_user.id, db)
    artifacts = (
        db.query(ProjectArtifact)
        .filter(ProjectArtifact.project_id == project_id, ProjectArtifact.content.isnot(None))
        .order_by(ProjectArtifact.created_at.asc())
        .all()
    )

    if not artifacts:
        raise HTTPException(status_code=404, detail="No artifacts available for download.")

    buffer = io.BytesIO()
    manifest = []
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for artifact in artifacts:
            safe_name = re.sub(r"[^a-zA-Z0-9_-]", "_", artifact.name).strip("_") or "artifact"
            if artifact.file_type == "playwright_project":
                files = (artifact.artifact_metadata or {}).get("files") or {}
                for path, file_content in files.items():
                    zf.writestr(f"{safe_name}/{path}", file_content or "")
                filename = f"{safe_name}/"
            else:
                filename = f"{safe_name}.docx"
                zf.writestr(filename, build_docx_bytes(artifact.name, artifact.content or ""))
            manifest.append({
                "id": artifact.id,
                "filename": filename,
                "artifact_type": artifact.artifact_type,
                "name": artifact.name,
                "generated_by": artifact.generated_by,
                "created_at": artifact.created_at.isoformat() if artifact.created_at else None,
            })
        zf.writestr("manifest.json", json.dumps(manifest, indent=2))

    buffer.seek(0)
    zip_name = re.sub(r"[^a-zA-Z0-9_-]", "_", project.name).strip("_") or "project"
    return StreamingResponse(
        buffer,
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename={zip_name}_artifacts.zip"},
    )